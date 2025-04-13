import base64
import io
import json
import os
import re
import zipfile
from datetime import datetime
from io import BytesIO
from os import PathLike
from typing import Any, BinaryIO

import fitz
import matplotlib.pyplot as plt
import torch
from docx import Document
from PIL import Image
from PyPDF2 import PdfReader
from torchvision import transforms
from transformers import AutoModel, AutoTokenizer

from jb_onboarding.constants import DOCS

FileInput = str | bytes | BinaryIO

fitz.TOOLS.mupdf_display_errors(False)


def transform_data(data: dict) -> dict:
    """
    Recursively traverse the data structure and replace any dictionary that
    has exactly the keys 'options' and 'selected' with the value from 'selected'.
    """
    if isinstance(data, dict):
        # If the dict has exactly the keys 'options' and 'selected', replace it
        if set(data.keys()) == {"options", "selected"}:
            if len(data["selected"]) == 1:
                return data["selected"][0]
            return data["selected"]
        else:
            # Otherwise, recursively transform its contents.
            return {key: transform_data(value) for key, value in data.items()}
    elif isinstance(data, list):
        # Process each element in the list.
        return [transform_data(item) for item in data]
    else:
        # Base case: if not a dict or list, return the data as-is.
        return data


def extract_contact_methods(contact_field: dict) -> dict:
    """
    Extracts telephone and email information from a dictionary with keys
    'Communication Medium' and None, where:
      - The 'Communication Medium' field contains a string like "Telephone +358 042 161 84 72"
      - The None key contains a string like "E-Mail andrea.fernandez@finet.fi"

    Returns a dictionary with two keys: "Telephone" and "E-mail".
    """
    new_fields = {}

    # Extract telephone information
    telephone_str = contact_field.get("Communication Medium", "")
    # Look for the word "Telephone" followed by any characters (the phone number)
    tel_match = re.search(r"Telephone\s+(.+)", telephone_str)
    if tel_match:
        # Remove any extra whitespace and use it
        new_fields["Telephone"] = tel_match.group(1).strip()

    # Extract email information
    email_str = contact_field.get(None, "")
    # Look for the word "E-Mail" followed by any characters (the email address)
    email_match = re.search(r"E-Mail\s+(.+)", email_str)
    if email_match:
        new_fields["E-mail"] = email_match.group(1).strip()

    return new_fields


def parse_address(addr_str):
    """
    Parses an address string of the form "StreetName StreetNumber, PostalCode City"
    and returns a dictionary with keys: city, street name, street number, postal code.
    """
    address = {"city": "", "street name": "", "street number": "", "postal code": ""}
    if not addr_str:
        return address
    parts = addr_str.split(",")
    if len(parts) >= 2:
        street_part = parts[0].strip()
        postal_city_part = parts[1].strip()
    else:
        street_part = addr_str
        postal_city_part = ""
    # Assume the last token in the street part is the street number.
    tokens = street_part.split()
    if tokens:
        if tokens[-1].isdigit():
            address["street number"] = tokens[-1]
            address["street name"] = " ".join(tokens[:-1])
        else:
            address["street name"] = street_part
    # For postal_city_part, assume the first token is postal code and the rest is city.
    tokens = postal_city_part.split()
    if tokens:
        address["postal code"] = tokens[0]
        address["city"] = " ".join(tokens[1:]) if len(tokens) > 1 else ""
    return address


def extract_contact_methods(contact_field: dict) -> dict:
    """
    Extract telephone and email information from a dictionary.
    The dictionary is expected to have a key "Communication Medium" with a string like:
      "Telephone +358 042 161 84 72"
    and a key None with a string like:
      "E-Mail andrea.fernandez@finet.fi"
    Returns a dict with "Telephone" and "E-mail" keys.
    """
    new_fields = {}
    telephone_str = contact_field.get("Communication Medium", "")
    tel_match = re.search(r"Telephone\s+(.+)", telephone_str)
    if tel_match:
        new_fields["Telephone"] = tel_match.group(1).strip()
    email_str = contact_field.get(None, "")
    email_match = re.search(r"E-Mail\s+(.+)", email_str)
    if email_match:
        new_fields["E-mail"] = email_match.group(1).strip()
    return new_fields


def extract_selected(field):
    """
    If field is a dict with keys 'options' and 'selected', return the first selected value.
    Otherwise return the field as is.
    """
    if isinstance(field, dict) and "selected" in field:
        selected = field.get("selected", [])
        if selected:
            return selected[0]
    return field


def extract_employment_history_manual(business_data: dict) -> list:
    """
    Extracts employment history records including company names for employed/self-employed individuals.
    """
    known_statuses = [
        "employee since",
        "self-employed since",
        "currently not employed since",
        "retired since",
        "homemaker/housewife since",
        "student since",
        "diplomat since",
        "military representative since",
        "other since",
    ]

    sections = []
    if "Current employment and function" in business_data:
        sections.append(business_data["Current employment and function"])
    if None in business_data:
        sections.append(business_data[None])

    records = []
    current_record = None

    for section in sections:
        for item in section:
            text = ""
            if isinstance(item, dict) and "options" in item:
                if item.get("selected"):
                    text = item["selected"][0] if item["selected"] else ""
                else:
                    text = item["options"][0] if item["options"] else ""
            elif isinstance(item, str):
                text = item
            else:
                continue

            text = text.strip().lower()

            # Handle previous professions
            if text.startswith("previous profession:"):
                profession = text[len("previous profession:") :].strip()
                if current_record:
                    current_record["previous_profession"] = profession
                continue

            # Handle employment statuses
            status_found = False
            for known in known_statuses:
                if text.startswith(known):
                    pattern = re.compile(re.escape(known) + r"(?:\s+(\d{4}))?")
                    m = pattern.match(text)
                    year = int(m.group(1)) if m and m.group(1) else None
                    status_label = known.replace(" since", "")

                    if current_record:
                        records.append(current_record)
                    current_record = {"status": status_label, "since": year, "previous_profession": "", "company": ""}
                    status_found = True
                    break

            if status_found:
                continue  # Skip further processing for this item

            # Extract company names using regex
            company_match = re.match(r"^name employer[: ]*(.*)", text)
            if company_match:
                company = company_match.group(1).strip()
                if current_record:
                    current_record["company"] = company
                continue

            company_match = re.match(r"^company name[: ]*(.*)", text)
            if company_match:
                company = company_match.group(1).strip()
                if current_record:
                    current_record["company"] = company

    if current_record:
        records.append(current_record)

    return records


def extract_wealth_and_assets(wealth_data: dict) -> dict:
    """
    Extracts wealth and asset information from the wealth_data dictionary.

    Expects wealth_data to have:
      - "Total wealth estimated": a list with a dictionary including a 'selected' list,
      - "Estimated Assets": a list with a dictionary including a 'selected' list.

    Returns a dictionary with keys:
      "wealth": the selected total wealth value (as a string),
      "assets": a list of selected assets.
    """
    result = {}

    # Extract "wealth" from "Total wealth estimated"
    wealth = None
    total_wealth = wealth_data.get("Total wealth estimated", [])
    if isinstance(total_wealth, list) and total_wealth:
        item = total_wealth[0]
        if isinstance(item, dict):
            selected = item.get("selected", [])
            if selected:
                wealth = selected[0]
    result["wealth"] = wealth

    # Extract "assets" from "Estimated Assets"
    assets = []
    estimated_assets = wealth_data.get("Estimated Assets", [])
    if isinstance(estimated_assets, list) and estimated_assets:
        item = estimated_assets[0]
        if isinstance(item, dict):
            selected = item.get("selected", [])
            if selected:
                assets = selected
    result["assets"] = assets

    # extract origins of wealth
    key = list(filter(lambda x: x.startswith("Origin of wealth "), list(wealth_data.keys())))[0]
    pattern = re.compile(r"(☒|☐)\s*([A-Za-z ]+)")
    origins = []

    for tick, label in pattern.findall(key):
        if tick == "☒":
            origins.append(label.strip().lower())
    result["origins of wealth"] = origins
    return result


def transform_profile(raw):
    new_profile = {}
    client_info = raw.get("Client Information", {})
    contact_info = raw.get("Account Holder – Contact Management and Services – Contact Info", {})
    personal_info = raw.get("Account Holder – Personal Info", {})
    account_info = raw.get("Account Information", {})
    aum_info = raw.get("Account Information – Asset Under Management", {})
    bussiness_info = raw.get("Account Holder – Professional and Economic Background – Business", {})
    # 1. Name: combine first/middle and last names.
    first = client_info.get("First/ Middle Name (s)", "").strip()
    last = client_info.get("Last Name", "").strip()
    new_profile["name"] = f"{first} {last}".strip()

    # 2. Address: parse the address string.
    addr_str = client_info.get("Address", "")
    new_profile["address"] = parse_address(addr_str)

    # 3. Country, birth date, nationality, passport numbers, and dates.
    new_profile["country_of_domicile"] = client_info.get("Country of Domicile", "").strip()
    new_profile["birth_date"] = client_info.get("Date of birth", "").strip()
    new_profile["nationality"] = client_info.get("Nationality", "").strip()
    new_profile["passport_number"] = client_info.get("Passport No/ Unique ID", "").strip()
    new_profile["passport_issue_date"] = client_info.get("ID Issue Date", "").strip()
    new_profile["passport_expiry_date"] = client_info.get("ID Expiry Date", "").strip()

    # 4. Gender: extract the selected value and map to a short string if applicable.
    gender_val = client_info.get("Gender")
    gender = extract_selected(gender_val)
    if isinstance(gender, str):
        # Map "Female" -> "F" and "Male" -> "M" (or keep as is)
        new_profile["gender"] = (
            "F" if gender.lower().startswith("f") else "M" if gender.lower().startswith("m") else gender
        )
    else:
        new_profile["gender"] = gender

    # 5. Phone number and email address from contact info.
    phone_email = extract_contact_methods(contact_info)
    new_profile["phone_number"] = phone_email.get("Telephone", "").strip()
    new_profile["email_address"] = phone_email.get("E-mail", "").strip()

    # 6. Marital Status:
    marital = personal_info.get("Marital Status")
    if isinstance(marital, list) and len(marital) > 0:
        new_profile["marital_status"] = extract_selected(marital[0])
    else:
        new_profile["marital_status"] = ""

    # 7. Secondary school (example from Education History):
    # Here we do a very naive parsing: if the education history string contains a year in parentheses,
    # we assume it's in the form "School Name (Year)".
    education_history = personal_info.get("Education History", "")
    school = {}
    if isinstance(education_history, str):
        m = re.search(r"(.+)\((\d{4})\)", education_history)
        if m:
            school["name"] = m.group(1).strip()
            school["graduation_year"] = int(m.group(2))
    school_type = personal_info.get("Highest education attained", "").lower()
    school["type"] = school_type

    new_profile["higher_education"] = school

    # 8. Higher education and employment history: we default to empty lists (or add your parsing).
    tmp = list(
        filter(
            lambda x: x["since"] is not None,
            extract_employment_history_manual(bussiness_info),
        )
    )
    new_profile["employment_background"] = tmp[0] if len(tmp) > 0 else {}

    # 9. AUM: for example, we map "Total Asset Under Management" (as savings) if available.
    new_profile["aum"] = {
        "savings": float(aum_info.get("Total Asset Under Management", 0))
        if str(aum_info.get("Total Asset Under Management", "0")).isdigit()
        else aum_info.get("Total Asset Under Management", 0),
    }

    # 10. Inheritance and real estate details: defaults.
    new_profile["inheritance_details"] = {}
    new_profile["real_estate_details"] = []

    # 11. Investment fields: extract using the helper.
    new_profile["investment_risk_profile"] = extract_selected(account_info.get("Investment Risk Profile", ""))
    new_profile["investment_horizon"] = extract_selected(account_info.get("Investment Horizon", ""))
    new_profile["investment_experience"] = extract_selected(account_info.get("Investment Experience", ""))
    new_profile["type_of_mandate"] = extract_selected(account_info.get("Type of Mandate", ""))
    # Preferred Markets: split comma-separated values.
    pm = account_info.get("Preferred Markets", "")
    if isinstance(pm, str):
        new_profile["preferred_markets"] = [p.strip() for p in pm.split(",")]
    else:
        new_profile["preferred_markets"] = pm

    # 12. Currency: from top-level field.
    new_profile["currency"] = raw.get("currency", "")

    # **Integration of Wealth and Assets**
    wealth_info = raw.get("Account Holder – Professional and Economic Background – Wealth", {})
    wealth_and_assets = extract_wealth_and_assets(wealth_info)
    new_profile["wealth"] = wealth_and_assets.get("wealth")
    new_profile["assets"] = wealth_and_assets.get("assets")
    new_profile["origins of wealth"] = wealth_and_assets.get("origins of wealth")

    return new_profile


class PassportParser:
    def __init__(self, model_id: str = "OpenGVLab/InternVL2_5-1B"):
        """
        Initializes the PassportParser instance by loading the specified model
        and its corresponding tokenizer. This instance is intended for OCR
        tasks on passport images using the InternVL2_5 series.

        Args:
            model_id (str): Hugging Face model identifier for the InternVL2_5 variant.
                            Defaults to "OpenGVLab/InternVL2_5-1B".
        """
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.model = (
            AutoModel.from_pretrained(
                model_id,
                trust_remote_code=True,
                torch_dtype=torch.float32,  # Keep or adjust as desired
                low_cpu_mem_usage=False,
            )
            .to(self.device)
            .eval()
        )
        self.tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)

    def __call__(self, path_to_passport: PathLike) -> tuple[dict, bool]:
        """
        Processes the passport image by loading and preprocessing it,
        then runs OCR inference to transcribe all the text visible in the image.
        """
        # Open the image and convert it to RGB.
        image = Image.open(path_to_passport).convert("RGB")
        passport_data = {}

        # Define the transformations: resize to 448x448, convert to tensor, and normalize.
        transform = transforms.Compose(
            [
                transforms.Resize((448, 448)),
                transforms.ToTensor(),
                transforms.Normalize(mean=(0.485, 0.456, 0.406), std=(0.229, 0.224, 0.225)),
            ]
        )

        # Apply the transformations and add a batch dimension.
        image_tensor = transform(image)
        pixel_values = image_tensor.unsqueeze(0).to(self.device)  # Shape: [1, 3, 448, 448]

        # Prepare the OCR prompt using the special <image> token.
        prompt = (
            "<image>\n"
            "Please extract the text from the passport image and output the information in a valid JSON object with exactly the following keys:\n"
            "  - Country\n"
            "  - Surname\n"
            "  - Code\n"
            "  - Passport Number\n"
            "  - Given Name\n"
            "  - Birth Date\n"
            "  - Citizenship\n"
            "  - Sex\n"
            "  - Issue Date\n"
            "  - Expiry Date\n"
            "  - MRZ line 1\n"
            "  - MRZ line 2\n"
            "For any field that is missing in the image, use an empty string as the value. Do not include any extra keys or text outside of the JSON."
        )
        output = None
        flag = True
        try:
            # Use the model's chat method to generate the OCR transcription.
            output = self.model.chat(
                self.tokenizer, pixel_values, prompt, generation_config={"max_new_tokens": 512, "do_sample": False}
            )

            # Decode the output and convert it to a JSON object.
            passport_data = json.loads(output[8:-3])
            flag = False
        except Exception as e:
            print(f"Error during OCR processing: {e}, {output}")
            # Use the model's chat method to generate the OCR transcription.

            try:
                output = self.model.chat(
                    self.tokenizer, pixel_values, prompt, generation_config={"max_new_tokens": 512, "do_sample": False}
                )

                # Decode the output and convert it to a JSON object.
                passport_data = json.loads(output[8:-3])
                flag = False
            except Exception as e:
                print(f"Second error during OCR processing: {e}, {output}")

        # get the signature
        signature = image.crop((239, 210, 359, 242)).tobytes()
        passport_data["Signature"] = base64.b64encode(signature).decode("utf-8")

        return passport_data, flag


class Preprocessor:
    def __init__(self, *args, **kwargs):
        """
        Initializes the Preprocessor instance. This class is responsible for
        extracting and processing data from a zip file containing various
        documents related to a user's profile.

        Args:
            **args: Additional positional arguments.
            **kwargs: Additional keyword arguments.
        """
        self.args = args
        self.kwargs = kwargs
        self.passport_parser = PassportParser()

    def __call__(
        self,
        path_to_zip: str,
        passport: FileInput | None = None,
        description: FileInput | None = None,
        account: FileInput | None = None,
        profile: FileInput | None = None,
    ) -> tuple[dict[str, Any], bool]:
        # Determine how to treat the zip archive itself.
        if path_to_zip is not None:
            # If we have a file path, derive temporary extraction dir from the name.
            temp_dir = str(path_to_zip).replace(".zip", "")
            with zipfile.ZipFile(path_to_zip, "r") as zip_ref:
                zip_ref.extractall(temp_dir)
                names = zip_ref.namelist()
                # profile.docx
                if "profile.docx" not in names:
                    raise FileNotFoundError("profile.docx not found in the zip archive.")
                with zip_ref.open("profile.docx") as profile_file:
                    doc = Document(profile_file)
                    profile_json = self._parse_profile_docx_to_json(doc)

                # passport.png
                if "passport.png" not in names:
                    raise FileNotFoundError("passport.png not found in the zip archive.")
                passport_data, flag = self.passport_parser(os.path.join(temp_dir, "passport.png"))

                # description.txt
                if "description.txt" not in names:
                    raise FileNotFoundError("description.txt not found in the zip archive.")
                with zip_ref.open("description.txt") as desc_file:
                    raw_desc = desc_file.read().decode("utf-8")
                    description_json = self._parse_description(raw_desc)

                # account.pdf
                if "account.pdf" not in names:
                    raise FileNotFoundError("account.pdf not found in the zip archive.")
                with zip_ref.open("account.pdf") as pdf_file:
                    account_pdf_bytes = pdf_file.read()
            account_parsed = self._extract_form_data_and_signature(account_pdf_bytes)
            output = {
                "profile": profile_json,
                "passport": passport_data,
                "description": description_json,
                "account": account_parsed,
            }
            return output, flag
        else:
            # Process profile.docx
            if isinstance(profile, bytes):
                profile_file = io.BytesIO(profile)
            elif hasattr(profile, "read"):
                profile_file = profile
            else:  # assume a file path
                profile_file = open(profile, "rb")
            doc = Document(profile_file)
            profile_json = self._parse_profile_docx_to_json(doc)
            # Close file if we opened it ourselves
            if not hasattr(profile, "read"):
                profile_file.close()

            # passport.png
            if isinstance(passport, bytes):
                passport_file = io.BytesIO(passport)
            elif hasattr(passport, "read"):
                passport_file = passport
            else:
                passport_file = open(passport, "rb")
            passport_data, flag = self.passport_parser(passport_file)
            if not hasattr(passport, "read"):
                passport_file.close()

            # Process description.txt

            if isinstance(description, bytes):
                raw_desc = description.decode("utf-8")
            elif hasattr(description, "read"):
                raw_desc = description.read().decode("utf-8")
            else:
                with open(description, "rb") as f:
                    raw_desc = f.read().decode("utf-8")
            description_json = self._parse_description(raw_desc)

            # Process account.pdf

            if isinstance(account, bytes):
                account_pdf_bytes = account
            elif hasattr(account, "read"):
                account_pdf_bytes = account.read()
            else:
                with open(account, "rb") as f:
                    account_pdf_bytes = f.read()

        account_parsed = self._extract_form_data_and_signature(account_pdf_bytes)

        output: dict[str, Any] = {
            "profile": transform_profile(profile_json),
            "passport": passport_data,
            "description": description_json,
            "account": account_parsed,
        }
        return output, flag

    def _parse_tick_field(self, text):
        pattern = r"(☒|☐)\s*([^☒☐]+)"
        matches = re.findall(pattern, text)
        options = []
        selected = []
        for mark, option in matches:
            option = option.strip()
            options.append(option)
            if mark == "☒":
                selected.append(option)
        return {"options": options, "selected": selected}

    def _clean_text(self, text):
        return " ".join(text.replace("\t", " ").split())

    def _parse_profile_docx_to_json(self, doc):
        result = {}
        current_section = "Default"
        result[current_section] = {}
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_values = [self._clean_text(cell.text) for cell in row.cells]
                if any(row_values):
                    table_data.append(row_values)
            if len(table_data) == 1 and len(table_data[0]) == 1:
                section_title = table_data[0][0]
                if section_title:
                    current_section = section_title
                    result[current_section] = {}
                continue
            for row in table_data:
                key = row[0] if row[0] else None
                # For the value, lookthrough the remaining cells for non-empty content.
                # (Sometimes there might be multiple pieces of information.)
                value_candidates = [cell for cell in row[1:] if cell]
                inline_value = None
                if not value_candidates and row[0]:
                    inline_kv_pattern = r"(\bE-?Mail\b|Telephone)\s+([\w@.\-+ ]+)"
                    inline_matches = re.findall(inline_kv_pattern, row[0])
                    for k, v in inline_matches:
                        key = k.strip()
                        inline_value = v.strip()
                        result[current_section][key] = inline_value
                    continue  # Skip further processing since we’ve already stored it
                if not value_candidates:
                    value = None
                elif len(value_candidates) == 1:
                    value = value_candidates[0]
                else:
                    value = value_candidates
                if value and isinstance(value, str) and ("☒" in value or "☐" in value):
                    value = self._parse_tick_field(value)
                elif isinstance(value, list):
                    new_vals = []
                    for v in value:
                        if "☒" in v or "☐" in v:
                            new_vals.append(self._parse_tick_field(v))
                        else:
                            new_vals.append(v)
                    value = new_vals
                if key in result[current_section]:
                    if not isinstance(result[current_section][key], list):
                        result[current_section][key] = [result[current_section][key]]
                    result[current_section][key].append(value)
                else:
                    result[current_section][key] = value
        return result

    def _parse_description(self, text: str) -> dict:
        lines = text.splitlines()
        result = {}
        current_key = None
        current_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_key is not None and current_lines and current_lines[-1] != "":
                    current_lines.append("")
                continue
            if stripped.endswith(":"):
                if current_key is not None:
                    result[current_key] = "\n".join(current_lines).strip()
                current_key = stripped[:-1].strip()
                current_lines = []
            else:
                if current_key is not None:
                    current_lines.append(stripped)
        if current_key is not None:
            result[current_key] = "\n".join(current_lines).strip()
        return result

    def _extract_form_fields(self, pdf_bytes):
        pdf_file = BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)
        fields = reader.get_fields()
        form_data = {}
        if fields:
            for field_name, field_info in fields.items():
                value = field_info.get("/V", "")
                form_data[field_name] = value
        else:
            print("No AcroForm fields found or the PDF is not a standard fillable form.")
        return form_data

    def _extract_signature_by_coordinates(self, pdf_bytes, ref_rect=(71, 582, 248, 625), dpi=300):
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count < 1:
            print("PDF has no pages.")
            return None
        page = doc.load_page(0)
        actual_rect = page.rect
        ref_a4_width, ref_a4_height = 595, 842
        scale_x = actual_rect.width / ref_a4_width
        scale_y = actual_rect.height / ref_a4_height
        ref_x0, ref_y0, ref_x1, ref_y1 = ref_rect
        target_rect = fitz.Rect(ref_x0 * scale_x, ref_y0 * scale_y, ref_x1 * scale_x, ref_y1 * scale_y)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        try:
            pix = page.get_pixmap(matrix=mat, clip=target_rect)
            # pix.save("signature1.png")
            png_bytes = pix.tobytes("png")
            return base64.b64encode(png_bytes).decode("utf-8")

        except Exception as e:
            print(f"Error extracting signature region: {e}")
            return None

    def _extract_form_data_and_signature(self, pdf_bytes):
        form_data = self._extract_form_fields(pdf_bytes)
        signature_b64 = self._extract_signature_by_coordinates(pdf_bytes)
        form_data["signature"] = signature_b64
        return form_data
